from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

def export_to_docx(content, filename="insights.docx"):
    doc = Document()
    doc.add_heading("AI Insights Dashboard", 0)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    doc.save(filename)
    return filename


def export_to_pdf(content, filename="insights.pdf"):
    doc = SimpleDocTemplate(filename)
    styles = getSampleStyleSheet()

    story = []
    for line in content.split("\n"):
        story.append(Paragraph(line, styles["Normal"]))

    doc.build(story)
    return filename